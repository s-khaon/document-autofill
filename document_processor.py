"""
Document Processor Module

This module handles the automation of document generation, including reading data from Excel,
processing Word templates, and downloading signature images from WeChat Drive.
"""

import os
import re
from typing import Optional, Dict, Any
import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from playwright.sync_api import sync_playwright
from tqdm import tqdm


class Config:  # pylint: disable=too-few-public-methods
    """Configuration constants for the application."""
    INPUT_DIR = '/Users/kang.song/projects/own/document-autofill/input'
    TEMPLATE_PATH = '/Users/kang.song/projects/own/document-autofill/模板文档2.docx'
    WECHAT_DRIVE_URL = 'https://drive.weixin.qq.com/s?k=ABMALgfzAA02hkJ1y6Ae8AMAYUACw'
    OUTPUT_BASE_DIR = 'output'


class WeChatDrive:
    """Handles interaction with WeChat Drive using Playwright."""

    def __init__(self):
        self.playwright = None
        self.browser = None
        self.context = None
        self.page = None

    def start(self):
        """Starts the Playwright browser session."""
        self.playwright = sync_playwright().start()
        # Headless=False for QR scan
        self.browser = self.playwright.chromium.launch(headless=False)
        self.context = self.browser.new_context()
        self.page = self.context.new_page()

    def login(self):
        """Navigates to the login page and waits for user login."""
        print("Opening WeChat Drive login page...")
        self.page.goto(Config.WECHAT_DRIVE_URL)

        # Wait for user to scan QR code
        print("Please scan the QR code to login...")
        try:
            self.page.wait_for_selector(
                '.user-avatar', timeout=60000)  # Wait 60s for login
            print("Login successful!")
        except Exception:  # pylint: disable=broad-exception-caught
            print("Login timeout or failed. Please try again.")
            input("Press Enter after you have successfully logged in...")

    def download_image(self, url: str, save_path: str) -> bool:
        """Downloads an image from the given URL to the save path."""
        try:
            self.page.goto(url)

            # Use Playwright's expect_download to handle the download event
            # We assume clicking the "下载" (Download) button triggers the download
            try:
                # Wait for the download button to be visible first
                download_btn = self.page.get_by_text("下载").first
                download_btn.wait_for(state="visible", timeout=10000)

                with self.page.expect_download(timeout=30000) as download_info:
                    download_btn.click()

                download = download_info.value
                download.save_as(save_path)
                return True
            except Exception as e:  # pylint: disable=broad-exception-caught
                print(f"Failed to download from {url}: {e}")
                return False

        except Exception as e:  # pylint: disable=broad-exception-caught
            print(f"Error navigating to {url}: {e}")
            return False

    def close(self):
        """Closes the browser and stops Playwright."""
        if self.browser:
            self.browser.close()
        if self.playwright:
            self.playwright.stop()


class DocumentProcessor:
    """Processes documents by reading Excel data and filling Word templates."""

    def __init__(self):
        self.drive = WeChatDrive()

    def process_date(self, date_val: Any, format_str: str) -> str:
        """Formats a date value according to the specified format string."""
        if pd.isna(date_val):
            return ""
        try:
            if isinstance(date_val, str):
                date_obj = pd.to_datetime(date_val)
            else:
                date_obj = date_val

            # Python format string mapping
            py_format = format_str.replace('yyyy', '%Y').replace(
                'mm', '%m').replace('dd', '%d')
            return date_obj.strftime(py_format)
        except Exception as e:  # pylint: disable=broad-exception-caught
            print(f"Date parsing error: {e}")
            return str(date_val)

    def extract_url(self, text: Any) -> str:
        """Extracts the first URL from the text."""
        if pd.isna(text):
            return ""
        text = str(text)
        # Regex to find URLs (simple version)
        url_match = re.search(r'(https?://[^\s]+)', text)
        if url_match:
            return url_match.group(0)
        return text

    def replace_text_in_paragraph(self, paragraph, data: Dict[str, Any]):
        """Replaces placeholders in a paragraph with data values while preserving formatting."""
        # Helper function to replace text in runs
        def replace_in_runs(para, placeholder, replacement):
            if placeholder not in para.text:
                return

            replaced = False
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
                    replaced = True

            if replaced:
                return

            # If the placeholder is split across runs, we fall back to paragraph text replacement
            # but try to preserve the style of the first run.
            style = para.runs[0].style if para.runs else None
            para.text = para.text.replace(placeholder, replacement)
            if style and para.runs:
                para.runs[0].style = style

        # 1. Date format: {{$date({{}}, "yyyy年mm月dd日")}}
        date_pattern = r'\{\{\$date\(\{\{(.*?)\}\},\s*"(.*?)"\)\}\}'
        for key, fmt in re.findall(date_pattern, paragraph.text):
            full_placeholder = f'{{{{$date({{{{{key}}}}}, "{fmt}")}}}}'
            if key in data:
                date_val = data.get(key)
                formatted_date = self.process_date(date_val, fmt)
                replace_in_runs(paragraph, full_placeholder, formatted_date)

        # 2. URL extraction: {{$url({{key}})}}
        url_pattern = r'\{\{\$url\(\{\{(.*?)\}\}\)\}\}'
        for key in re.findall(url_pattern, paragraph.text):
            full_placeholder = f'{{{{$url({{{{{key}}}}})}}}}'
            if key in data:
                raw_text = data.get(key)
                url = self.extract_url(raw_text)
                replace_in_runs(paragraph, full_placeholder, url)

        # 3. {{期限}} -> 可接受授权时间 (Legacy support)
        if '{{期限}}' in paragraph.text:
            replace_in_runs(paragraph, '{{期限}}', str(data.get('可接受授权时间', '')))

        # 4. Regular placeholders
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                replace_in_runs(paragraph, placeholder, str(value))

    def insert_image_at_placeholder(self, doc, placeholder: str,
                                    image_path: str, width_inches: float = 1.5) -> bool:
        """Inserts an image at the specified placeholder location."""
        # Search in paragraphs
        if self._search_and_replace_image(doc.paragraphs, placeholder, image_path, width_inches):
            return True

        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if self._search_and_replace_image(
                            cell.paragraphs, placeholder, image_path, width_inches):
                        return True
        return False

    def _search_and_replace_image(self, paragraphs, placeholder, image_path, width_inches):
        """Helper to search and replace image in a list of paragraphs."""
        for para in paragraphs:
            if placeholder in para.text:
                # The placeholder might be {{签名图片}} or {签名图片} or similar.
                # If we found it, we need to replace it with the image.

                # Check if it's the only content (common case for signature cells)
                if para.text.strip() == placeholder:
                    para.text = ""  # Clear text
                    run = para.add_run()
                    run.add_picture(image_path, width=Inches(width_inches))
                    return True

                # If mixed content, replacing text is tricky because image is appended.
                # But for signature, it's usually acceptable to append.
                # NOTE: para.text assignment destroys formatting and runs!
                # We should replace text in runs instead if possible.

                # Try to find the run containing the placeholder
                replaced = False
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, "")
                        run.add_picture(image_path, width=Inches(width_inches))
                        replaced = True
                        return True

                if not replaced:
                    # If split across runs, we have to fall back to para.text assignment
                    # This might lose formatting of other text in the same paragraph
                    para.text = para.text.replace(placeholder, "")
                    run = para.add_run()
                    run.add_picture(image_path, width=Inches(width_inches))
                    return True
        return False

    def get_excel_files(self) -> list[str]:
        """Returns a list of all Excel files in the input directory."""
        if not os.path.exists(Config.INPUT_DIR):
            print(f"Input directory not found: {Config.INPUT_DIR}")
            return []

        files = []
        for f in os.listdir(Config.INPUT_DIR):
            if f.endswith('.xlsx') and not f.startswith('~$'):
                files.append(os.path.join(Config.INPUT_DIR, f))
        return files

    def get_excel_hyperlinks(self, excel_path: str) -> Dict[int, Optional[str]]:
        """Extracts hyperlinks from the '请签本名' column in the Excel file."""
        wb = load_workbook(excel_path)
        ws = wb.active

        header_row = [cell.value for cell in ws[1]]
        try:
            link_col_idx = header_row.index("请签本名") + 1
        except ValueError:
            print(
                f"Column '请签本名' not found in {os.path.basename(excel_path)}!")
            return {}

        links = {}
        for i, row in enumerate(ws.iter_rows(min_row=2), start=0):
            cell = row[link_col_idx-1]
            if cell.hyperlink:
                links[i] = cell.hyperlink.target
            else:
                links[i] = None
        return links

    def process_single_row(self, index: int, row, links: Dict[int, Optional[str]]):
        """Processes a single row of data."""
        try:
            data = row.to_dict()
            nickname = data.get('您的平台昵称')
            if pd.isna(nickname):
                nickname = f'User_{index}'

            output_dir = self._prepare_output_dir(nickname)
            doc = self._create_filled_document(data)

            # Handle Image
            link = links.get(index)
            self._handle_signature_image(link, output_dir, doc, nickname)

            # Save document
            doc_path = os.path.join(output_dir, f"{nickname}_卡赫视频授权书.docx")
            doc.save(doc_path)
        except Exception as e:  # pylint: disable=broad-exception-caught
            print(f"Error processing row {index}: {e}")

    def _prepare_output_dir(self, nickname: str) -> str:
        """Creates and returns the output directory path."""
        folder_name = f"[{nickname}]x卡赫视频授权"
        output_dir = os.path.join(Config.OUTPUT_BASE_DIR, folder_name)
        os.makedirs(output_dir, exist_ok=True)
        return output_dir

    def _create_filled_document(self, data: Dict[str, Any]) -> Document:
        """Creates a new document and fills it with data."""
        doc = Document(Config.TEMPLATE_PATH)

        # Replace placeholders
        for para in doc.paragraphs:
            self.replace_text_in_paragraph(para, data)

        for table in doc.tables:
            for r in table.rows:
                for cell in r.cells:
                    for para in cell.paragraphs:
                        self.replace_text_in_paragraph(para, data)
        return doc

    def _handle_signature_image(self, link: Optional[str],
                                output_dir: str, doc: Document, nickname: str):
        """Downloads and inserts the signature image if available."""
        if link:
            image_path = os.path.join(output_dir, 'signature.png')
            if self.drive.download_image(link, image_path):
                if not self.insert_image_at_placeholder(doc, "{{签名图片}}", image_path):
                    print(
                        f"Warning: Placeholder {{签名图片}} not found for {nickname}")
            else:
                print(f"Failed to download image for {nickname}")

    def run(self):
        """Main execution method."""
        excel_files = self.get_excel_files()
        if not excel_files:
            print("No Excel files found in input directory.")
            return

        self.drive.start()
        self.drive.login()

        for excel_path in excel_files:
            print(f"Processing Excel file: {os.path.basename(excel_path)}...")
            try:
                df = pd.read_excel(excel_path)
                links = self.get_excel_hyperlinks(excel_path)

                print("Processing rows...")
                for index, row in tqdm(df.iterrows(), total=len(df)):
                    self.process_single_row(index, row, links)
            except Exception as e:  # pylint: disable=broad-exception-caught
                print(f"Error processing file {excel_path}: {e}")

        self.drive.close()
        print("Processing complete!")


if __name__ == "__main__":
    processor = DocumentProcessor()
    processor.run()
