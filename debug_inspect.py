import pandas as pd
from docx import Document
import os

excel_path = '/Users/kang.song/projects/own/document-autofill/input/视频授权书.xlsx'
word_path = '/Users/kang.song/projects/own/document-autofill/模板文档2.docx'

def inspect_data():
    if os.path.exists(excel_path):
        df = pd.read_excel(excel_path)
        print("Excel Columns:", df.columns.tolist())
        # Check the first row data, especially the video link
        first_row = df.iloc[0].to_dict()
        print("First row data:", first_row)
        if '授权视频链接' in first_row:
             print("Link value type:", type(first_row['授权视频链接']))
             print("Link value:", first_row['授权视频链接'])

    if os.path.exists(word_path):
        doc = Document(word_path)
        print("\n--- Word Template Structure ---")
        for i, para in enumerate(doc.paragraphs):
            if '{{' in para.text:
                print(f"Para {i}: {para.text}")
                for j, run in enumerate(para.runs):
                    print(f"  Run {j}: '{run.text}'")

        for i, table in enumerate(doc.tables):
            print(f"Table {i}:")
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        if '{{' in para.text:
                            print(f"  Row {r_idx} Col {c_idx} Para {p_idx}: {para.text}")
                            for j, run in enumerate(para.runs):
                                print(f"    Run {j}: '{run.text}'")

if __name__ == "__main__":
    inspect_data()
